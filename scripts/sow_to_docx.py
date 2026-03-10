"""Convert sales-productivity-sow.md to a polished .docx file matching the CRM Wizards Project Outline format."""
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    if level == 1:
        h.font.size = Pt(18)
    elif level == 2:
        h.font.size = Pt(14)
    else:
        h.font.size = Pt(12)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders)
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C" w:val="clear"/>')
            cell._tc.get_or_add_tcPr().append(shading)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
                    run.font.size = Pt(9.5)
    for i, row in enumerate(table.rows):
        if i == 0:
            continue
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
        if i % 2 == 0:
            for cell in row.cells:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F6FA" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)


def add_rich_paragraph(doc_or_section, text, style_name=None, bold=False):
    """Add a paragraph with inline bold (**text**) support."""
    p = doc_or_section.add_paragraph(style=style_name)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


def add_table_from_rows(doc, header, rows):
    """Add a formatted table."""
    cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h.strip().replace('**', ''))
        run.bold = True
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            clean = val.strip().replace('**', '')
            run = cell.paragraphs[0].add_run(clean)
            if val.strip().startswith('**'):
                run.bold = True
    style_table(table)
    doc.add_paragraph()


def parse_md_table(lines, start):
    """Parse a markdown table starting at line index `start`."""
    header_line = lines[start].strip()
    cols = [c.strip() for c in header_line.strip('|').split('|')]
    idx = start + 2
    rows = []
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        row = [c.strip() for c in lines[idx].strip().strip('|').split('|')]
        rows.append(row)
        idx += 1
    return cols, rows, idx


# ---- Read the markdown ----
with open('/home/elikiedrowski12/ParmM_Zoho/sales-productivity-sow.md') as f:
    md = f.read()

lines = md.split('\n')
i = 0

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip horizontal rules
    if stripped == '---':
        i += 1
        continue

    # Skip empty lines
    if stripped == '':
        i += 1
        continue

    # Headings
    if stripped.startswith('### '):
        text = stripped[4:].strip()
        doc.add_heading(text, level=3)
        i += 1
        continue
    if stripped.startswith('## '):
        text = stripped[3:].strip()
        doc.add_heading(text, level=2)
        i += 1
        continue
    if stripped.startswith('# '):
        text = stripped[2:].strip()
        doc.add_heading(text, level=1)
        i += 1
        continue

    # Tables
    if stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
        cols, rows, end = parse_md_table(lines, i)
        add_table_from_rows(doc, cols, rows)
        i = end
        continue

    # Code blocks
    if stripped.startswith('```'):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        i += 1
        p = doc.add_paragraph()
        run = p.add_run('\n'.join(code_lines))
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
        p._p.get_or_add_pPr().append(shading)
        continue

    # Bullet points (indented sub-bullets)
    if line.startswith('  - ') or line.startswith('    - '):
        text = stripped[2:]
        text = text.replace('~~', '')
        p = add_rich_paragraph(doc, text)
        p.style = doc.styles['List Bullet']
        indent_level = (len(line) - len(line.lstrip())) // 2
        p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent_level)
        i += 1
        continue

    # Bullet points (top-level)
    if stripped.startswith('- '):
        text = stripped[2:]
        text = text.replace('~~', '')
        p = add_rich_paragraph(doc, text)
        p.style = doc.styles['List Bullet']
        i += 1
        continue

    # Numbered list
    m = re.match(r'^(\d+)\.\s+', stripped)
    if m:
        text = stripped[m.end():]
        text = text.replace('~~', '')
        p = add_rich_paragraph(doc, text)
        p.style = doc.styles['List Number']
        i += 1
        continue

    # Regular paragraph
    text = stripped.replace('~~', '')
    if text:
        add_rich_paragraph(doc, text)
    i += 1

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

output = '/home/elikiedrowski12/ParmM_Zoho/sales-productivity-sow.docx'
doc.save(output)
print(f'Saved to {output}')
